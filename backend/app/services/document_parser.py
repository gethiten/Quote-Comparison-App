"""Parse carrier quote documents (PDF, DOCX, XLSX) and extract structured data.

Uses Azure AI Content Understanding first when available; falls back to the
OpenAI extraction agent and then simple rule-based parsing.
"""
import base64
import io
import logging
import time

import requests
from azure.identity import DefaultAzureCredential

from app.config import settings
from app.schemas.schemas import QuoteBase

logger = logging.getLogger(__name__)


def _has_value(value: object) -> bool:
    if value in (None, "", [], {}):
        return False
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return False
        return any(ch.isalnum() for ch in stripped)
    return True


def _needs_enrichment(result: QuoteBase | None) -> bool:
    if result is None:
        return True
    important_fields = ("carrier_name", "expiry_date", "underwriting_notes")
    return any(not _has_value(getattr(result, field, None)) for field in important_fields)


def _merge_results(primary: QuoteBase | None, secondary: QuoteBase | None) -> QuoteBase | None:
    if primary is None:
        return secondary
    if secondary is None:
        return primary

    merged = secondary.model_dump()
    for key, value in primary.model_dump().items():
        if _has_value(value):
            merged[key] = value
    return QuoteBase(**merged)


async def parse_quote_document(
    file_data: bytes, filename: str, content_type: str
) -> QuoteBase | None:
    """Extract structured quote data from a file."""
    # Start with the custom Content Understanding analyzer.
    cu_result = _extract_with_content_understanding(file_data, filename, content_type)
    if cu_result and not _needs_enrichment(cu_result):
        return cu_result

    text = _extract_text(file_data, filename, content_type)
    if not text:
        return cu_result

    ai_result = None
    try:
        from app.agents.quote_extraction_agent import extract_quote_with_ai

        ai_result = await extract_quote_with_ai(text, filename)
    except Exception:
        logger.exception("OpenAI text extraction failed for %s", filename)

    merged = _merge_results(cu_result, ai_result)

    # Always let the rule-based parser fill any remaining structured gaps.
    return _merge_results(merged, _rule_based_parse(text))


def _get_access_token() -> str:
    credential = DefaultAzureCredential()
    token = credential.get_token("https://cognitiveservices.azure.com/.default")
    return token.token


def _extract_with_content_understanding(
    file_data: bytes, filename: str, content_type: str
) -> QuoteBase | None:
    endpoint = getattr(settings, "AZURE_CONTENT_UNDERSTANDING_ENDPOINT", "") or settings.AZURE_OPENAI_ENDPOINT
    analyzer_id = getattr(settings, "AZURE_CONTENT_UNDERSTANDING_ANALYZER_ID", "insuranceQuoteExtractor")
    api_version = "2025-11-01"

    if not endpoint:
        return None

    try:
        token = _get_access_token()
        analyze_url = (
            f"{endpoint.rstrip('/')}/contentunderstanding/analyzers/"
            f"{analyzer_id}:analyze?api-version={api_version}"
        )
        body = {
            "inputs": [
                {
                    "data": base64.b64encode(file_data).decode("ascii"),
                    "mimeType": _guess_mime_type(filename, content_type),
                }
            ]
        }
        resp = requests.post(
            analyze_url,
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
            },
            json=body,
            timeout=20,
        )
        if resp.status_code != 202:
            logger.info(
                "Content Understanding unavailable for %s: HTTP %s",
                filename,
                resp.status_code,
            )
            return None

        result_url = resp.headers.get("Operation-Location")
        if not result_url:
            return None

        for _ in range(20):
            time.sleep(1)
            poll = requests.get(
                result_url,
                headers={"Authorization": f"Bearer {token}"},
                timeout=20,
            )
            if poll.status_code != 200:
                continue
            payload = poll.json()
            status = payload.get("status")
            if status == "Succeeded":
                extracted = _parse_cu_result(payload)
                mapped = _map_cu_fields_to_quote(extracted, filename)
                if any(value is not None for key, value in mapped.items() if key not in {"raw_file_url", "source_filename"}):
                    return QuoteBase(**mapped)
                return None
            if status in {"Failed", "Canceled"}:
                return None
    except Exception:
        logger.exception("Content Understanding extraction failed for %s", filename)

    return None


def _guess_mime_type(filename: str, content_type: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf") or "pdf" in content_type:
        return "application/pdf"
    if lower.endswith(".docx") or "wordprocessingml" in content_type:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if lower.endswith(".xlsx") or "spreadsheetml" in content_type:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return content_type or "application/octet-stream"


def _parse_cu_result(result: dict) -> dict:
    extracted: dict[str, object] = {}
    contents = result.get("result", {}).get("contents", [])
    for content in contents:
        fields = content.get("fields", {})
        for field_name, field_data in fields.items():
            if not isinstance(field_data, dict):
                continue
            field_type = field_data.get("type", "string")
            if field_type == "number":
                value = field_data.get("valueNumber")
            elif field_type == "date":
                value = field_data.get("valueDate")
            elif field_type == "boolean":
                value = field_data.get("valueBoolean")
            else:
                value = field_data.get("valueString")
            if value is not None:
                extracted[field_name] = value
    return extracted


def _map_cu_fields_to_quote(extracted: dict, filename: str) -> dict:
    bi_period = _to_float(extracted.get("BIPeriodMonths"))
    return {
        "carrier_name": extracted.get("CarrierName"),
        "quote_number": extracted.get("QuoteNumber"),
        "quote_date": extracted.get("QuoteDate"),
        "effective_date": extracted.get("EffectiveDate"),
        "expiry_date": extracted.get("ExpiryDate"),
        "building_limit": _to_float(extracted.get("BuildingLimit")),
        "valuation_basis": _normalize_valuation_basis(extracted.get("ValuationBasis")),
        "coverage_form": _normalize_coverage_form(extracted.get("CoverageForm")),
        "coinsurance": int(_to_float(extracted.get("Coinsurance"))) if _to_float(extracted.get("Coinsurance")) is not None else None,
        "bpp_limit": _to_float(extracted.get("BPPLimit")),
        "business_interruption_limit": _to_float(extracted.get("BusinessInterruptionLimit")),
        "bi_period_months": int(bi_period) if bi_period is not None else None,
        "gl_per_occurrence": _to_float(extracted.get("GLPerOccurrence")),
        "gl_aggregate": _to_float(extracted.get("GLAggregate")),
        "aop_deductible": _to_float(extracted.get("AOPDeductible")),
        "wind_hail_deductible_pct": _to_float(extracted.get("WindHailDeductiblePct")),
        "flood_limit": _to_float(extracted.get("FloodLimit")),
        "earthquake_limit": _to_float(extracted.get("EarthquakeLimit")),
        "equipment_breakdown": _to_bool(extracted.get("EquipmentBreakdown")),
        "ordinance_or_law": _to_bool(extracted.get("OrdinanceOrLaw")),
        "annual_premium": _to_float(extracted.get("AnnualPremium")),
        "underwriting_notes": extracted.get("UnderwritingNotes"),
        "raw_file_url": None,
        "source_filename": filename,
    }


def _to_float(val: object) -> float | None:
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    cleaned = str(val).replace(",", "").replace("$", "").strip()
    try:
        return float(cleaned)
    except ValueError:
        return None


def _to_bool(val: object) -> bool | None:
    if val is None:
        return None
    if isinstance(val, bool):
        return val
    normalized = str(val).strip().lower()
    if normalized in {"yes", "true", "included", "incl", "y"}:
        return True
    if normalized in {"no", "false", "excluded", "excl", "n"}:
        return False
    return None


def _normalize_valuation_basis(val: object) -> str | None:
    if val is None:
        return None
    normalized = str(val).strip().lower()
    if not normalized:
        return None
    if "replacement" in normalized or normalized == "rc":
        return "RC"
    if "actual cash" in normalized or normalized == "acv":
        return "ACV"
    return None


def _normalize_coverage_form(val: object) -> str | None:
    if val is None:
        return None
    normalized = str(val).strip().lower()
    if not normalized:
        return None
    if "special" in normalized:
        return "Special"
    if "broad" in normalized:
        return "Broad"
    if "basic" in normalized:
        return "Basic"
    return None


def _extract_text(data: bytes, filename: str, content_type: str) -> str:
    """Extract raw text content from a file."""
    lower = filename.lower()

    if lower.endswith(".pdf") or "pdf" in content_type:
        return _extract_pdf(data)
    elif lower.endswith(".docx") or "wordprocessingml" in content_type:
        return _extract_docx(data)
    elif lower.endswith(".xlsx") or "spreadsheetml" in content_type:
        return _extract_xlsx(data)
    return ""

def _extract_pdf(data: bytes) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) if c is not None else "" for c in row]
            if any(vals):
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def _rule_based_parse(text: str) -> QuoteBase | None:
    """Rule-based extraction used to fill any remaining gaps."""
    import re
    from datetime import datetime

    def _find_number(*patterns: str) -> float | None:
        for pattern in patterns:
            m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if m:
                val = m.group(1).replace(",", "").replace("$", "").strip()
                try:
                    return float(val)
                except ValueError:
                    continue
        return None

    def _find_str(*patterns: str) -> str | None:
        for pattern in patterns:
            m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if m:
                return m.group(1).strip()
        return None

    def _find_date(*patterns: str) -> str | None:
        for pattern in patterns:
            m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if not m:
                continue
            raw = re.sub(r"\s+", " ", m.group(1).strip().strip("|:-"))
            for fmt in ("%B %d, %Y", "%b %d, %Y", "%m/%d/%Y", "%m-%d-%Y"):
                try:
                    return datetime.strptime(raw, fmt).date().isoformat()
                except ValueError:
                    pass
        return None

    def _extract_notes() -> str | None:
        section_patterns = (
            r"RISK NOTES\s*(?:[|:]\s*)?(.+?)(?=CONDITIONS\s*&\s*EXCLUSIONS|TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"CONDITIONS\s*&\s*EXCLUSIONS\s*(?:[|:]\s*)?(.+?)(?=TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"UNDERWRITING NOTES\s*(?:[|:]\s*)?(.+?)(?=TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"IMPORTANT CONDITIONS AND UNDERWRITING NOTES\s*(.+?)(?:This document is a quotation only|THE HARTFORD|$)",
            r"IMPORTANT NOTES\s*&\s*CONDITIONS\s*(.+?)(?:Current FEMA flood determination|consent of AIG|$)",
        )
        chunks: list[str] = []
        for pattern in section_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL):
                raw_chunk = match.group(1)
                cleaned_lines: list[str] = []
                for line in raw_chunk.splitlines():
                    clean = re.sub(r"\s*\|\s*", " ", line).strip(" |-:\t")
                    if len(clean) < 6 or not any(ch.isalpha() for ch in clean):
                        continue
                    if re.fullmatch(r"[A-Z][A-Z\s/&-]+", clean):
                        continue
                    cleaned_lines.append(clean)
                candidate = " ".join(cleaned_lines).strip()
                if candidate and candidate not in chunks:
                    chunks.append(candidate)
        return " ".join(chunks) if chunks else None

    lowered = text.lower()
    carrier_name = None
    if "travelers" in lowered:
        carrier_name = "Travelers"
    elif "hartford" in lowered:
        carrier_name = "The Hartford"
    elif re.search(r"\baig\b", lowered):
        carrier_name = "AIG"
    elif "zurich" in lowered:
        carrier_name = "Zurich"

    quote_number = _find_str(
        r"quote\s*(?:number|#|no\.?)\s*(?:[|:]\s*|\s+)([A-Z0-9][A-Z0-9\-/]+)"
    )
    quote_date = _find_date(
        r"quote\s*date\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"^\s*([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
    )
    effective_date = _find_date(
        r"coverage effective\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"policy effective\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"policy period[:\s]*([A-Z][a-z]+\s+\d{1,2},\s+\d{4})\s+(?:to|\|)",
        r"effective\s*(?:date)?\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
    )
    expiry_date = _find_date(
        r"policy period[:\s]*[A-Z][a-z]+\s+\d{1,2},\s+\d{4}\s+(?:to|\|)\s+([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"quote expires(?: at [\d: AMPamp]+)?(?:\s+on)?\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"quote expiration\s*(?:date)?\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"this quote expires\s*(?:on\s*)?([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"quote is valid through ([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
    )

    premium = _find_number(
        r"TOTAL ANNUAL PREMIUM\s*[|:]\s*\$?([\d,]+)",
        r"(?:total\s+(?:annual\s+)?premium|annual\s+premium)[:\s]*\$?([\d,]+)",
    )
    building = _find_number(
        r"Insured Value\s*[—-]\s*Building\s*[|:]\s*\$?([\d,]+)",
        r"Building\s*[|:]\s*\$?([\d,]+)\s*[|]",
        r"(?:building\s+limit|building)[:\s]*\$?([\d,]+)",
    )
    deductible = _find_number(r"All Other Perils\s*[|:]\s*\$?([\d,]+)", r"(?:all\s+other\s+perils|aop)[:\s]*\$?([\d,]+)")
    bpp = _find_number(r"Business Personal Property\s*[|:]\s*\$?([\d,]+)", r"(?:business\s+personal\s+property|bpp)[:\s]*\$?([\d,]+)")
    bi = _find_number(
        r"Business Income\s*/\s*Extra Expense\s*[|:]\s*\$?([\d,]+)",
        r"Loss of Rents\s*[|:]\s*\$?([\d,]+)",
        r"(?:business\s+(?:income|interruption))[^$]*\$?([\d,]+)",
    )
    bi_period = _find_number(
        r"Business Income\s*/\s*Extra Expense\s*[|:]\s*\$?[\d,]+\s*[|]\s*(\d+)-Month",
        r"Loss of Rents\s*[|:]\s*(?:Actual Loss Sustained\s*[|:]\s*)?(\d+)\s*-?\s*months?",
    )
    flood = None if re.search(r"flood.+(?:excluded|not included|not offered)", text, re.IGNORECASE) else _find_number(r"flood[^$]*\$?([\d,]+)")
    earthquake = None if re.search(r"earthquake.+(?:excluded|not included|not offered)", text, re.IGNORECASE) else _find_number(r"earthquake[^$]*\$?([\d,]+)")
    wind_hail = _find_number(r"(\d+(?:\.\d+)?)%\s+(?:of\s+TIV\s+)?(?:per-occurrence\s+)?deductible applies for Wind/Hail", r"wind\s*/\s*hail deductible\D+(\d+(?:\.\d+)?)%")

    valuation = None
    if re.search(r"actual\s+cash\s+value|\bACV\b", text, re.IGNORECASE):
        valuation = "ACV"
    elif re.search(r"replacement\s+cost", text, re.IGNORECASE):
        valuation = "RC"

    coverage_form = None
    if re.search(r"special\s+form|causes\s+of\s+loss\s+form\s*[|:]\s*special", text, re.IGNORECASE):
        coverage_form = "Special"
    elif re.search(r"broad\s+form|causes\s+of\s+loss\s+form\s*[|:]\s*broad", text, re.IGNORECASE):
        coverage_form = "Broad"
    elif re.search(r"basic\s+form|causes\s+of\s+loss\s+form\s*[|:]\s*basic", text, re.IGNORECASE):
        coverage_form = "Basic"

    coinsurance = _find_number(r"(\d{1,3})%\s+coinsurance")
    equipment_breakdown = False if re.search(r"equipment breakdown.+(?:not included|excluded)", text, re.IGNORECASE) else True if re.search(r"equipment breakdown.+included", text, re.IGNORECASE) else None
    ordinance_or_law = False if re.search(r"ordinance\s+or\s+law.+(?:not included|excluded)", text, re.IGNORECASE) else True if re.search(r"ordinance\s+or\s+law.+included", text, re.IGNORECASE) else None
    underwriting_notes = _extract_notes() or _find_str(
        r"NOTES\s*(.+?)SECTION 3",
        r"(This quote expires.+?)(?:Current FEMA flood determination|consent of AIG|$)",
        r"review\. This quote is valid through .+?Please read your",
        r"underwriting notes[:\s]*(.+)$",
    )

    if not premium and not building and not quote_number:
        return None

    return QuoteBase(
        carrier_name=carrier_name,
        quote_number=quote_number,
        quote_date=quote_date,
        effective_date=effective_date,
        expiry_date=expiry_date,
        annual_premium=premium,
        building_limit=building,
        aop_deductible=deductible,
        bpp_limit=bpp,
        business_interruption_limit=bi,
        bi_period_months=int(bi_period) if bi_period is not None else None,
        flood_limit=flood,
        earthquake_limit=earthquake,
        valuation_basis=valuation,
        coverage_form=coverage_form,
        coinsurance=int(coinsurance) if coinsurance is not None else None,
        wind_hail_deductible_pct=wind_hail,
        equipment_breakdown=equipment_breakdown,
        ordinance_or_law=ordinance_or_law,
        underwriting_notes=underwriting_notes,
    )
