"""Azure Function: Blob trigger to process uploaded quote documents.

When a file is uploaded to the 'quote-documents' container, this function:
1. Sends the blob to Azure AI Content Understanding (insuranceQuoteExtractor custom analyzer).
2. Maps the structured extracted fields to the database schema.
3. Inserts a new Quote record into Azure PostgreSQL.
"""

import json
import logging
import os
import re
import time
import base64
import io

import azure.functions as func
import psycopg2
import requests
from azure.identity import DefaultAzureCredential

logger = logging.getLogger(__name__)

CU_ENDPOINT = os.environ.get("CU_ENDPOINT", "")
CU_API_VERSION = "2025-11-01"
ANALYZER_ID = "insuranceQuoteExtractor"
STORAGE_ACCOUNT = os.environ.get("STORAGE_ACCOUNT", "quotecomparestr2026")
OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT", CU_ENDPOINT)
OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4.1")
OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2025-01-01-preview")

ACS_CONNECTION_STRING = os.environ.get("ACS_CONNECTION_STRING", "")
NOTIFICATION_EMAIL = os.environ.get("NOTIFICATION_EMAIL", "")
ACS_SENDER = os.environ.get("ACS_SENDER", "")


def main(blob: func.InputStream):
    filename = blob.name or ""
    logger.info("Processing uploaded file: %s (%d bytes)", filename, blob.length or 0)

    file_data = blob.read()
    if not file_data:
        logger.warning("Empty blob, skipping: %s", filename)
        return

    logger.info("Storage-trigger extraction pipeline v2 active for %s", filename)

    # --- Step 1: Extract with Content Understanding, then fall back to text parsing ---
    extracted = extract_with_content_understanding(file_data, filename)

    if extracted and _has_structured_fields(extracted):
        quote_data = map_fields(extracted, filename)
        logger.info("Mapped quote data from Content Understanding: %s", json.dumps(quote_data, default=str))
    else:
        if extracted:
            logger.warning(
                "Content Understanding returned no structured fields for %s; falling back to text extraction",
                filename,
            )
        text = (extracted or {}).get("_full_text") or extract_text(file_data, filename)
        if not text:
            logger.warning("No data extracted from %s", filename)
            return

        fallback_extracted = extract_with_openai_text(text, filename) or rule_based_text_parse(text)
        if not fallback_extracted:
            logger.warning("No data extracted from %s", filename)
            return

        quote_data = map_fields(fallback_extracted, filename)
        logger.info("Mapped quote data from text fallback: %s", json.dumps(quote_data, default=str))

    # --- Step 3: Insert into PostgreSQL ---
    insert_result = insert_quote(quote_data, filename)
    if not insert_result:
        logger.warning("Quote was extracted from %s but no database row was inserted", filename)
        return
    if insert_result.get("skipped"):
        logger.warning(
            "Skipped duplicate quote from %s (existing quote_id=%s)",
            filename,
            insert_result.get("quote_id"),
        )
        return

    logger.info(
        "Successfully processed and inserted quote from %s (quote_id=%s)",
        filename,
        insert_result.get("quote_id"),
    )

    # --- Step 4: Send email notification ---
    send_notification(quote_data, filename)


def _get_access_token() -> str:
    """Get an Azure AD token using managed identity."""
    credential = DefaultAzureCredential()
    token = credential.get_token("https://cognitiveservices.azure.com/.default")
    return token.token


def extract_with_content_understanding(
    file_data: bytes, filename: str
) -> dict | None:
    """Use Azure AI Content Understanding prebuilt-document to extract data."""
    if not CU_ENDPOINT:
        logger.error("CU_ENDPOINT not configured")
        return None

    try:
        token = _get_access_token()
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        }

        # Send file content as base64 (avoids storage RBAC issues)
        file_b64 = base64.b64encode(file_data).decode("ascii")

        # Determine MIME type
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            mime = "application/pdf"
        elif lower_name.endswith(".docx"):
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif lower_name.endswith(".xlsx"):
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        else:
            mime = "application/octet-stream"

        # Start analysis using Content Understanding
        analyze_url = (
            f"{CU_ENDPOINT.rstrip('/')}/contentunderstanding/analyzers/"
            f"{ANALYZER_ID}:analyze?api-version={CU_API_VERSION}"
        )
        body = {"inputs": [{"data": file_b64, "mimeType": mime}]}

        resp = requests.post(analyze_url, headers=headers, json=body, timeout=30)
        if resp.status_code != 202:
            logger.error(
                "Content Understanding analyze failed: HTTP %d - %s",
                resp.status_code,
                resp.text,
            )
            return None

        # Poll for result
        result_url = resp.headers.get("Operation-Location")
        if not result_url:
            logger.error("No Operation-Location header in response")
            return None

        for _ in range(60):  # poll up to 2 minutes
            time.sleep(2)
            poll_resp = requests.get(
                result_url,
                headers={"Authorization": f"Bearer {token}"},
                timeout=30,
            )
            if poll_resp.status_code != 200:
                logger.warning("Poll returned HTTP %d", poll_resp.status_code)
                continue

            result = poll_resp.json()
            status = result.get("status", "")
            if status == "Succeeded":
                return _parse_cu_result(result)
            if status in ("Failed", "Canceled"):
                logger.error("Analysis %s: %s", status, json.dumps(result))
                return None

        logger.error("Content Understanding analysis timed out for %s", filename)
        return None

    except Exception:
        logger.exception("Content Understanding extraction failed for %s", filename)
        return None


def _parse_cu_result(result: dict) -> dict:
    """Parse the Content Understanding result into a dict of extracted data."""
    extracted = {}
    contents = result.get("result", {}).get("contents", [])

    for content in contents:
        # Get markdown text (keep as fallback)
        markdown = content.get("markdown", "")
        extracted["_full_text"] = markdown

        # Get structured fields from custom analyzer
        fields = content.get("fields", {})
        for field_name, field_data in fields.items():
            if not isinstance(field_data, dict):
                continue
            field_type = field_data.get("type", "string")
            if field_type == "number":
                val = field_data.get("valueNumber")
            elif field_type == "date":
                val = field_data.get("valueDate")
            elif field_type == "boolean":
                val = field_data.get("valueBoolean")
            else:
                val = field_data.get("valueString")
            if val is not None:
                extracted[field_name] = val
                confidence = field_data.get("confidence")
                if confidence is not None:
                    logger.info("  %s = %s (confidence: %.3f)", field_name, val, confidence)
                else:
                    logger.info("  %s = %s", field_name, val)

    logger.info("Extracted %d fields from CU custom analyzer", len([k for k in extracted if not k.startswith("_")]))
    return extracted


def _has_structured_fields(extracted: dict) -> bool:
    return any(
        value not in (None, "", [], {})
        for key, value in extracted.items()
        if not str(key).startswith("_")
    )


def extract_text(file_data: bytes, filename: str) -> str:
    lower_name = filename.lower()
    try:
        if lower_name.endswith(".pdf"):
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(file_data))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n".join(parts)

        if lower_name.endswith(".docx"):
            from docx import Document

            doc = Document(io.BytesIO(file_data))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

        if lower_name.endswith(".xlsx"):
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
            parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if vals:
                        parts.append(" | ".join(vals))
            return "\n".join(parts)
    except Exception:
        logger.exception("Text extraction fallback failed for %s", filename)

    return ""


def extract_with_openai_text(text: str, filename: str) -> dict | None:
    if not OPENAI_ENDPOINT or not OPENAI_DEPLOYMENT or not text:
        return None

    system_prompt = """You are an expert commercial property insurance quote parser.
Return ONLY valid JSON with exactly these keys:
CarrierName, QuoteNumber, QuoteDate, EffectiveDate, ExpiryDate,
BuildingLimit, ValuationBasis, CoverageForm, Coinsurance,
BPPLimit, BusinessInterruptionLimit, BIPeriodMonths,
GLPerOccurrence, GLAggregate, AOPDeductible, WindHailDeductiblePct,
FloodLimit, EarthquakeLimit, EquipmentBreakdown, OrdinanceOrLaw,
AnnualPremium, UnderwritingNotes.

Rules:
- Use null when unknown.
- Monetary values must be plain numbers.
- ValuationBasis must be RC or ACV.
- CoverageForm must be Special, Broad, or Basic.
- WindHailDeductiblePct should be the numeric percent value (for example 3 for 3%).
- EquipmentBreakdown and OrdinanceOrLaw should be true/false/null.
- UnderwritingNotes should summarize key conditions, exclusions, and remarks.
"""

    try:
        token = _get_access_token()
        url = (
            f"{OPENAI_ENDPOINT.rstrip('/')}/openai/deployments/{OPENAI_DEPLOYMENT}"
            f"/chat/completions?api-version={OPENAI_API_VERSION}"
        )
        body = {
            "messages": [
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": f"Extract structured quote fields from this {filename} document:\n\n{text[:15000]}",
                },
            ],
            "temperature": 0,
            "max_tokens": 1800,
        }
        resp = requests.post(
            url,
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
            },
            json=body,
            timeout=60,
        )
        if resp.status_code != 200:
            logger.warning("OpenAI text fallback failed: HTTP %s - %s", resp.status_code, resp.text)
            return None

        payload = resp.json()
        content = payload["choices"][0]["message"]["content"].strip()
        if content.startswith("```"):
            content = content.split("\n", 1)[1]
            if content.endswith("```"):
                content = content[: content.rfind("```")]
        data = json.loads(content)
        return data if isinstance(data, dict) else None
    except Exception:
        logger.exception("OpenAI text fallback failed for %s", filename)
        return None


def rule_based_text_parse(text: str) -> dict | None:
    from datetime import datetime

    def find_number(*patterns: str):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if not match:
                continue
            val = match.group(1).replace(",", "").replace("$", "").strip()
            try:
                return float(val)
            except ValueError:
                continue
        return None

    def find_string(*patterns: str):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                return match.group(1).strip()
        return None

    def find_date(*patterns: str):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if not match:
                continue
            raw = re.sub(r"\s+", " ", match.group(1).strip().strip("|:-"))
            for fmt in ("%B %d, %Y", "%b %d, %Y", "%m/%d/%Y", "%m-%d-%Y"):
                try:
                    return datetime.strptime(raw, fmt).date().isoformat()
                except ValueError:
                    pass
        return None

    def find_notes():
        section_patterns = (
            r"RISK NOTES\s*(?:[|:]\s*)?(.+?)(?=CONDITIONS\s*&\s*EXCLUSIONS|TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"CONDITIONS\s*&\s*EXCLUSIONS\s*(?:[|:]\s*)?(.+?)(?=TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"UNDERWRITING NOTES\s*(?:[|:]\s*)?(.+?)(?=TOTAL ANNUAL PREMIUM|QUOTE\s+(?:NUMBER|DATE)|$)",
            r"IMPORTANT CONDITIONS AND UNDERWRITING NOTES\s*(.+?)(?:This document is a quotation only|THE HARTFORD|$)",
        )
        chunks = []
        for pattern in section_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL):
                cleaned_lines = []
                for line in match.group(1).splitlines():
                    clean = re.sub(r"\s*\|\s*", " ", line).strip(" |-:\t")
                    if len(clean) < 6 or not any(ch.isalpha() for ch in clean):
                        continue
                    if re.fullmatch(r"[A-Z][A-Z\s/&-]+", clean):
                        continue
                    cleaned_lines.append(clean)
                candidate = " ".join(cleaned_lines).strip()
                if candidate and candidate not in chunks:
                    chunks.append(candidate)
        return " ".join(chunks) if chunks else None

    lowered = text.lower()
    carrier_name = None
    if "hartford" in lowered:
        carrier_name = "The Hartford"
    elif "travelers" in lowered:
        carrier_name = "Travelers"
    elif re.search(r"\baig\b", lowered):
        carrier_name = "AIG"
    elif "zurich" in lowered:
        carrier_name = "Zurich Insurance Group"

    quote_date = find_date(
        r"quote\s*date\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"^\s*([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
    )
    effective_date = find_date(
        r"coverage effective\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"policy effective\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"Policy Period:\s*([A-Z][a-z]+\s+\d{1,2},\s+\d{4})\s+(?:to|\|)",
    )
    expiry_date = find_date(
        r"Policy Period:\s*[A-Z][a-z]+\s+\d{1,2},\s+\d{4}\s+(?:to|\|)\s+([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
        r"quote expires(?:\s+on)?\s*(?:[|:]\s*|\s+)([A-Z][a-z]+\s+\d{1,2},\s+\d{4})",
    )

    extracted = {
        "CarrierName": carrier_name,
        "QuoteNumber": find_string(
            r"Quote\s*#\s*([A-Z0-9\-]+)",
            r"quote\s*(?:number|#|no\.?)\s*(?:[|:]\s*|\s+)([A-Z0-9][A-Z0-9\-/]+)",
        ),
        "QuoteDate": quote_date,
        "EffectiveDate": effective_date,
        "ExpiryDate": expiry_date,
        "BuildingLimit": find_number(
            r"Insured Value\s*[—-]\s*Building\s*\|\s*\$?([\d,]+)",
            r"Building\s*\|\s*\$?([\d,]+)\s*\|",
            r"insured value of\s*\$?([\d,]+)",
        ),
        "Coinsurance": find_number(r"(\d{1,3})%\s+coinsurance"),
        "BPPLimit": find_number(
            r"Business Personal Property\s*\|\s*\$?([\d,]+)",
            r"(?:business\s+personal\s+property|bpp)[^\d$]*\$?([\d,]+)",
        ),
        "BusinessInterruptionLimit": find_number(
            r"Business Income\s*/\s*Extra Expense\s*\|\s*\$?([\d,]+)",
            r"(?:business\s+(?:income|interruption))[^\d$]*\$?([\d,]+)",
        ),
        "BIPeriodMonths": find_number(r"Business Income\s*/\s*Extra Expense\s*\|\s*\$?[\d,]+\s*\|\s*(\d+)-Month"),
        "GLPerOccurrence": find_number(r"General Liability \(Per Occ\.\)\s*\|\s*\$?([\d,]+)"),
        "GLAggregate": find_number(r"General Liability \(Aggregate\)\s*\|\s*\$?([\d,]+)"),
        "AOPDeductible": find_number(r"All Other Perils\s*\|\s*\$?([\d,]+)", r"(?:all\s+other\s+perils|aop)[^\d$]*\$?([\d,]+)"),
        "WindHailDeductiblePct": find_number(r"WIND\s*/\s*HAIL DEDUCTIBLE\s*A\s*(\d+(?:\.\d+)?)%"),
        "AnnualPremium": find_number(
            r"TOTAL ANNUAL PREMIUM\s*\|\s*\$?([\d,]+)",
            r"(?:annual\s+premium|total\s+premium)[^\d$]*\$?([\d,]+)",
        ),
        "UnderwritingNotes": find_notes() or find_string(r"IMPORTANT CONDITIONS AND UNDERWRITING NOTES\s*(.+?)This document is a quotation only"),
    }

    if re.search(r"actual\s+cash\s+value|\bACV\b", text, re.IGNORECASE):
        extracted["ValuationBasis"] = "ACV"
    elif re.search(r"replacement\s+cost", text, re.IGNORECASE):
        extracted["ValuationBasis"] = "RC"

    if re.search(r"broad\s+form", text, re.IGNORECASE):
        extracted["CoverageForm"] = "Broad"
    elif re.search(r"special\s+form", text, re.IGNORECASE):
        extracted["CoverageForm"] = "Special"
    elif re.search(r"basic\s+form", text, re.IGNORECASE):
        extracted["CoverageForm"] = "Basic"

    if re.search(r"equipment\s+breakdown.+not included", text, re.IGNORECASE):
        extracted["EquipmentBreakdown"] = False
    elif re.search(r"equipment\s+breakdown.+included", text, re.IGNORECASE):
        extracted["EquipmentBreakdown"] = True

    if re.search(r"ordinance\s+or\s+law.+not included|ordinance\s+or\s+law.+excluded", text, re.IGNORECASE):
        extracted["OrdinanceOrLaw"] = False
    elif re.search(r"ordinance\s+or\s+law.+included", text, re.IGNORECASE):
        extracted["OrdinanceOrLaw"] = True

    if re.search(r"flood.+not offered|flood.+excluded", text, re.IGNORECASE):
        extracted["FloodLimit"] = None
    if re.search(r"earthquake.+not offered|earthquake.+excluded", text, re.IGNORECASE):
        extracted["EarthquakeLimit"] = None

    return extracted if _has_structured_fields(extracted) else None


def map_fields(extracted: dict, filename: str) -> dict:
    """Map custom analyzer structured fields to the quote database fields."""

    def parse_bool(val):
        if val is None:
            return None
        if isinstance(val, bool):
            return val
        v = str(val).lower().strip()
        if v in ("yes", "true", "included", "incl", "y"):
            return True
        if v in ("no", "false", "excluded", "excl", "n"):
            return False
        if "included" in v or re.search(r"\d", v):
            return True
        return None

    def to_float(val):
        if val is None:
            return None
        if isinstance(val, (int, float)):
            return float(val)
        cleaned = re.sub(r"[,$\s]", "", str(val))
        try:
            return float(cleaned)
        except (ValueError, TypeError):
            return None

    def normalize_percent(val):
        num = to_float(val)
        if num is None:
            return None
        if 0 < abs(num) <= 1:
            return round(num * 100, 2)
        return num

    def normalize_valuation_basis(val):
        if val is None:
            return None
        v = str(val).lower()
        if "replacement" in v or v.strip() == "rc":
            return "RC"
        if "actual cash" in v or "acv" in v:
            return "ACV"
        return None

    def normalize_coverage_form(val):
        if val is None:
            return None
        v = str(val).lower()
        if "special" in v:
            return "Special"
        if "broad" in v:
            return "Broad"
        if "basic" in v:
            return "Basic"
        return None

    coinsurance = normalize_percent(extracted.get("Coinsurance"))
    bi_period = to_float(extracted.get("BIPeriodMonths"))

    return {
        "carrier_name": extracted.get("CarrierName"),
        "quote_number": extracted.get("QuoteNumber"),
        "quote_date": extracted.get("QuoteDate"),
        "effective_date": extracted.get("EffectiveDate"),
        "expiry_date": extracted.get("ExpiryDate"),
        "building_limit": to_float(extracted.get("BuildingLimit")),
        "valuation_basis": normalize_valuation_basis(extracted.get("ValuationBasis")),
        "coverage_form": normalize_coverage_form(extracted.get("CoverageForm")),
        "coinsurance": int(abs(coinsurance)) if coinsurance is not None else None,
        "bpp_limit": to_float(extracted.get("BPPLimit")),
        "business_interruption_limit": to_float(extracted.get("BusinessInterruptionLimit")),
        "bi_period_months": int(abs(bi_period)) if bi_period is not None else None,
        "gl_per_occurrence": to_float(extracted.get("GLPerOccurrence")),
        "gl_aggregate": to_float(extracted.get("GLAggregate")),
        "aop_deductible": to_float(extracted.get("AOPDeductible")),
        "wind_hail_deductible_pct": normalize_percent(extracted.get("WindHailDeductiblePct")),
        "flood_limit": to_float(extracted.get("FloodLimit")),
        "earthquake_limit": to_float(extracted.get("EarthquakeLimit")),
        "equipment_breakdown": parse_bool(extracted.get("EquipmentBreakdown")),
        "ordinance_or_law": parse_bool(extracted.get("OrdinanceOrLaw")),
        "annual_premium": to_float(extracted.get("AnnualPremium")),
        "underwriting_notes": extracted.get("UnderwritingNotes"),
        "source_filename": filename.split("/")[-1] if "/" in filename else filename,
        "raw_file_url": (
            f"https://{STORAGE_ACCOUNT}.blob.core.windows.net/{filename}"
        ),
    }


def _normalize_carrier_name(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = " ".join(str(value).strip().split())
    return normalized or None


def _normalize_quote_number(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = re.sub(r"\s+", " ", str(value).strip())
    return normalized or None


def _canonical_quote_number(value: str | None) -> str | None:
    normalized = _normalize_quote_number(value)
    if not normalized:
        return None
    canonical = re.sub(r"[^A-Za-z0-9]+", "", normalized).upper()
    return canonical or None


def _find_duplicate_quote_id(cur, carrier_id, quote_number: str | None):
    canonical_quote_number = _canonical_quote_number(quote_number)
    if not carrier_id or not canonical_quote_number:
        return None

    cur.execute(
        """
        SELECT quote_id
        FROM quotes
        WHERE carrier_id = %s
          AND UPPER(REGEXP_REPLACE(COALESCE(quote_number, ''), '[^A-Za-z0-9]+', '', 'g')) = %s
        LIMIT 1
        """,
        (carrier_id, canonical_quote_number),
    )
    row = cur.fetchone()
    return row[0] if row else None


def insert_quote(quote_data: dict, filename: str):
    """Insert extracted quote data into Azure PostgreSQL."""
    conn = psycopg2.connect(
        host=os.environ["PG_HOST"],
        database=os.environ["PG_DATABASE"],
        user=os.environ["PG_USER"],
        password=os.environ["PG_PASSWORD"],
        sslmode="require",
    )
    try:
        cur = conn.cursor()

        quote_data["carrier_name"] = _normalize_carrier_name(quote_data.get("carrier_name"))
        quote_data["quote_number"] = _normalize_quote_number(quote_data.get("quote_number"))

        # Resolve or create carrier
        carrier_id = _resolve_carrier(cur, quote_data.get("carrier_name"))

        duplicate_quote_id = _find_duplicate_quote_id(
            cur,
            carrier_id,
            quote_data.get("quote_number"),
        )
        if duplicate_quote_id:
            logger.warning(
                "Duplicate quote detected for %s: carrier=%s, quote_number=%s, existing_quote_id=%s",
                filename,
                quote_data.get("carrier_name"),
                _normalize_quote_number(quote_data.get("quote_number")),
                duplicate_quote_id,
            )
            return {
                "quote_id": str(duplicate_quote_id),
                "skipped": True,
                "reason": "duplicate_quote",
            }

        # Resolve or create a property/account so uploads work even in an empty database.
        property_id, account_id = _resolve_or_create_property(cur, quote_data)

        cur.execute(
            """
            INSERT INTO quotes (
                property_id, carrier_id, quote_number, quote_date,
                effective_date, expiry_date, building_limit, valuation_basis,
                coverage_form, coinsurance, bpp_limit,
                business_interruption_limit, bi_period_months,
                gl_per_occurrence, gl_aggregate, aop_deductible,
                wind_hail_deductible_pct, flood_limit, earthquake_limit,
                equipment_breakdown, ordinance_or_law, annual_premium,
                underwriting_notes, raw_file_url, source_filename
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s
            ) RETURNING quote_id
        """,
            (
                property_id,
                carrier_id,
                quote_data.get("quote_number"),
                quote_data.get("quote_date"),
                quote_data.get("effective_date"),
                quote_data.get("expiry_date"),
                quote_data.get("building_limit"),
                quote_data.get("valuation_basis"),
                quote_data.get("coverage_form"),
                quote_data.get("coinsurance"),
                quote_data.get("bpp_limit"),
                quote_data.get("business_interruption_limit"),
                quote_data.get("bi_period_months"),
                quote_data.get("gl_per_occurrence"),
                quote_data.get("gl_aggregate"),
                quote_data.get("aop_deductible"),
                quote_data.get("wind_hail_deductible_pct"),
                quote_data.get("flood_limit"),
                quote_data.get("earthquake_limit"),
                quote_data.get("equipment_breakdown"),
                quote_data.get("ordinance_or_law"),
                quote_data.get("annual_premium"),
                quote_data.get("underwriting_notes"),
                quote_data.get("raw_file_url"),
                quote_data.get("source_filename"),
            ),
        )
        quote_id = cur.fetchone()[0]

        comparison_id = _ensure_comparison(cur, account_id)
        cur.execute(
            "SELECT COALESCE(MAX(display_order), -1) + 1 FROM comparison_quotes WHERE comparison_id = %s",
            (comparison_id,),
        )
        display_order = cur.fetchone()[0]
        cur.execute(
            """
            INSERT INTO comparison_quotes (comparison_id, quote_id, display_order)
            VALUES (%s, %s, %s)
            """,
            (comparison_id, quote_id, display_order),
        )

        conn.commit()
        logger.info(
            "Quote inserted for property %s (comparison %s)", property_id, comparison_id
        )
        return {
            "quote_id": str(quote_id),
            "property_id": str(property_id),
            "comparison_id": str(comparison_id),
        }
    except Exception:
        conn.rollback()
        logger.exception("Failed to insert quote from %s", filename)
        raise
    finally:
        conn.close()


def _resolve_carrier(cur, carrier_name: str | None):
    """Find or create a carrier by name and return its ID."""
    carrier_name = _normalize_carrier_name(carrier_name)
    if not carrier_name:
        # Create an unknown carrier placeholder
        cur.execute(
            "INSERT INTO carriers (carrier_name) VALUES (%s) "
            "ON CONFLICT DO NOTHING RETURNING carrier_id",
            ("Unknown",),
        )
        row = cur.fetchone()
        if row:
            return row[0]
        cur.execute(
            "SELECT carrier_id FROM carriers WHERE carrier_name = %s",
            ("Unknown",),
        )
        return cur.fetchone()[0]

    cur.execute(
        "SELECT carrier_id FROM carriers WHERE LOWER(BTRIM(carrier_name)) = LOWER(%s)",
        (carrier_name,),
    )
    row = cur.fetchone()
    if row:
        return row[0]

    cur.execute(
        "INSERT INTO carriers (carrier_name) VALUES (%s) RETURNING carrier_id",
        (carrier_name,),
    )
    return cur.fetchone()[0]


def _resolve_or_create_property(cur, quote_data: dict):
    """Find an existing property or create a default account/property for uploads."""
    cur.execute("SELECT property_id, account_id FROM properties ORDER BY created_at LIMIT 1")
    row = cur.fetchone()
    if row:
        return row[0], row[1]

    account_name = "Storage Upload Imports"
    cur.execute(
        "SELECT account_id FROM accounts WHERE client_name = %s LIMIT 1",
        (account_name,),
    )
    row = cur.fetchone()
    if row:
        account_id = row[0]
    else:
        cur.execute(
            "INSERT INTO accounts (client_name, address) VALUES (%s, %s) RETURNING account_id",
            (account_name, "Auto-created from Azure Storage quote upload"),
        )
        account_id = cur.fetchone()[0]
        logger.info("Created default account %s for uploaded quotes", account_id)

    cur.execute(
        """
        INSERT INTO properties (
            account_id, type, sub_type, address, city, state, zip, insured_value
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING property_id
        """,
        (
            account_id,
            "office",
            "Imported Quote",
            quote_data.get("source_filename") or "Uploaded quote document",
            "Unknown",
            "NA",
            "00000",
            quote_data.get("building_limit") or 0,
        ),
    )
    property_id = cur.fetchone()[0]
    logger.warning(
        "No properties existed. Created default property %s for blob upload processing",
        property_id,
    )
    return property_id, account_id


def _ensure_comparison(cur, account_id):
    """Ensure there is a comparison row for the account so the UI can show the quote."""
    cur.execute(
        "SELECT comparison_id FROM comparisons WHERE account_id = %s ORDER BY created_at LIMIT 1",
        (account_id,),
    )
    row = cur.fetchone()
    if row:
        return row[0]

    cur.execute(
        """
        INSERT INTO comparisons (account_id, client_name, producer, notes, status)
        VALUES (%s, %s, %s, %s, %s)
        RETURNING comparison_id
        """,
        (
            account_id,
            "Storage Upload Imports",
            "System",
            "Auto-created from Azure Storage quote upload",
            "active",
        ),
    )
    comparison_id = cur.fetchone()[0]
    logger.info("Created comparison %s for account %s", comparison_id, account_id)
    return comparison_id


def send_notification(quote_data: dict, filename: str):
    """Send email notification after a quote is successfully processed."""
    missing_settings = [
        name
        for name, value in {
            "ACS_CONNECTION_STRING": ACS_CONNECTION_STRING,
            "NOTIFICATION_EMAIL": NOTIFICATION_EMAIL,
            "ACS_SENDER": ACS_SENDER,
        }.items()
        if not value
    ]
    if missing_settings:
        logger.warning(
            "Email notification skipped — missing app settings: %s",
            ", ".join(missing_settings),
        )
        return

    try:
        from azure.communication.email import EmailClient

        carrier = quote_data.get("carrier_name") or "Unknown"
        premium = quote_data.get("annual_premium")
        premium_str = f"${premium:,.0f}" if premium else "N/A"
        building = quote_data.get("building_limit")
        building_str = f"${building:,.0f}" if building else "N/A"
        source_file = quote_data.get("source_filename") or filename

        subject = f"New Quote Processed: {carrier}"
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; color: #333;">
            <h2 style="color: #1e40af;">Quote Processed Successfully</h2>
            <p>A new insurance quote has been extracted and saved to the database.</p>
            <table style="border-collapse: collapse; margin: 16px 0;">
                <tr><td style="padding: 6px 12px; font-weight: bold; border-bottom: 1px solid #e2e8f0;">Carrier</td>
                    <td style="padding: 6px 12px; border-bottom: 1px solid #e2e8f0;">{carrier}</td></tr>
                <tr><td style="padding: 6px 12px; font-weight: bold; border-bottom: 1px solid #e2e8f0;">Annual Premium</td>
                    <td style="padding: 6px 12px; border-bottom: 1px solid #e2e8f0;">{premium_str}</td></tr>
                <tr><td style="padding: 6px 12px; font-weight: bold; border-bottom: 1px solid #e2e8f0;">Building Limit</td>
                    <td style="padding: 6px 12px; border-bottom: 1px solid #e2e8f0;">{building_str}</td></tr>
                <tr><td style="padding: 6px 12px; font-weight: bold; border-bottom: 1px solid #e2e8f0;">Quote Number</td>
                    <td style="padding: 6px 12px; border-bottom: 1px solid #e2e8f0;">{quote_data.get("quote_number") or "N/A"}</td></tr>
                <tr><td style="padding: 6px 12px; font-weight: bold; border-bottom: 1px solid #e2e8f0;">Effective Date</td>
                    <td style="padding: 6px 12px; border-bottom: 1px solid #e2e8f0;">{quote_data.get("effective_date") or "N/A"}</td></tr>
                <tr><td style="padding: 6px 12px; font-weight: bold;">Source File</td>
                    <td style="padding: 6px 12px;">{source_file}</td></tr>
            </table>
            <p style="color: #64748b; font-size: 12px;">This is an automated notification from QuoteCompare Pro.</p>
        </body>
        </html>
        """

        client = EmailClient.from_connection_string(ACS_CONNECTION_STRING)
        message = {
            "senderAddress": ACS_SENDER,
            "recipients": {
                "to": [{"address": NOTIFICATION_EMAIL}],
            },
            "content": {
                "subject": subject,
                "html": html_body,
            },
        }

        poller = client.begin_send(message)
        result = poller.result()
        logger.info("Email notification sent (id: %s) to %s", result.get("id", ""), NOTIFICATION_EMAIL)

    except Exception:
        logger.exception("Failed to send email notification for %s", filename)
